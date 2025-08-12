##################################################
## testy.py
##################################################
## Author: clenahan@cloud2gnd.com
## Copyright: Copyright 2024
## Version: 1.0
##################################################

import docx.table
import docx.text
import docx.text.paragraph
from bcolors import bcolors
from wordhelper import getAcceptedText
from excelhelper import setHeaderCell

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT,WD_SECTION
from docx.enum.dml import MSO_COLOR_TYPE

import openpyxl
from openpyxl.cell.text import InlineFont
from openpyxl.formatting.rule import CellIsRule, FormulaRule
from openpyxl.cell.rich_text import TextBlock, CellRichText
import docx
import os
import re
import nltk

#import ssl
#try:
#    _create_unverified_https_context = ssl._create_unverified_context
#except AttributeError:
#    pass
#else:
#    ssl._create_default_https_context = _create_unverified_https_context



class HeadingTracking:

    def __init__(self) -> None:
        self.lastHeaderLevel = 0
        self.outlinelevel={}
        self.headers={}
        self.ignoreableSection = False
        self.testcase = False  # True if we are in a test case section, so we ignore all headers
        self.testprocedure = False  # Added for Test Procedure tracking
        pass

    def __incrementOrSet__(self,level):
        if level in self.outlinelevel:
            self.outlinelevel[level] += 1
        else:
            self.outlinelevel[level] = 1

    def pushHeader(self,style,text):
        level=style2level(style)
        if level > self.lastHeaderLevel:
            #indent
            self.__incrementOrSet__(level)
        elif level < self.lastHeaderLevel:
            #undent
            #reset all sub to zero
            self.__incrementOrSet__(level)
            for x in range(level+1,8):
                foo = self.outlinelevel.pop(x,"foo")
                bar = self.headers.pop(x,"bar")
                #self.outlinelevel[level] = 0
        else:
            #level is the same
            self.__incrementOrSet__(level)

        self.headers[level] = text
        self.lastHeaderLevel = level

    def getHeaders(self):
        return self.headers.copy()
    
    def getOutlineLevel(self):
        text = ""
        bFirst = True
        for x in range (1,7):
            if x in self.outlinelevel:
                text = text + "%s%s" % ("." if x > 1 else "" , self.outlinelevel[x])
        return text

################################################################################################################################################

def processTCA(filename):
    if  not os.path.exists(filename):
        print(bcolors.FAIL + "File '%s' does not exist" % filename + bcolors.ENDC)
        return
    
    #nltk.download('punkt')
    #nltk.download('punkt_tab')

    doc = Document(filename)

    reqlist = []

    parseDocument(doc, processNonHeader,lambda x,y,z: reqlist.append((x,y,z)))

    outputRequirements(filename,reqlist)


def assignOrInsert(cellRichText,index,value):
    if index >= len(cellRichText):
        cellRichText.insert(index,value)
    else:
        cellRichText[index]=value
    

def createRichText(text,targetstatements):
    ret = CellRichText(text)

    for (targetword,font) in targetstatements.items():
        regex="\W" + targetword +"\W"
        for (idx,block) in reversed(list(enumerate(ret))):
            if isinstance(block,str):
                lastpos = 0
                bFound = False
                originaltext = block
                currentIdx = idx
                for group in  re.finditer(regex,block):
                    bFound= True
                    #ret[currentIdx]=originaltext[lastpos:group.start()]
                    assignOrInsert(ret,currentIdx,originaltext[lastpos:group.start()])
                    currentIdx += 1
                    #ret.insert(currentIdx,TextBlock(font,originaltext[group.start():group.end()]))
                    assignOrInsert(ret,currentIdx,TextBlock(font,originaltext[group.start():group.end()]))
                    currentIdx += 1
                    lastpos = group.end()
                if bFound and lastpos < len(originaltext):
                    #ret.insert(currentIdx,originaltext[lastpos:])
                    assignOrInsert(ret,currentIdx,originaltext[lastpos:])

    return ret

        

def outputRequirements(filename,reqlist):
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.sheet_view.zoomScale = 140
    setHeaderCell(sheet,1,"Entry",5)
    setHeaderCell(sheet,2,"Section #",5)
    setHeaderCell(sheet,3,"Heading 1",25)
    setHeaderCell(sheet,4,"Heading 2",25)
    setHeaderCell(sheet,5,"Heading 3",25)
    setHeaderCell(sheet,6,"Name",15)
    setHeaderCell(sheet,7,"Requirement",100)
    setHeaderCell(sheet,8,"Obligation",15)
    setHeaderCell(sheet,9,"Test Case ID",15)
    setHeaderCell(sheet,10,"Comments/Notes",15)
    setHeaderCell(sheet,11,"Test  Type (C,B,F)",15)

    targetstatements={}
    targetstatements["shall"]=InlineFont(sheet.cell(row = 1, column =1).font)
    targetstatements["shall"].color = "00FF0000"
    targetstatements["may"]=InlineFont(sheet.cell(row = 1, column =1).font)
    targetstatements["may"].color = "0000FF00"

    nRow=2
    for (section,headerlist,requirement) in reqlist:
        sheet.cell(row = nRow, column =1).value = nRow -1
        sheet.cell(row = nRow, column =2).value = section
        nCol = 3
        for header in headerlist.values():
            sheet.cell(row = nRow, column =nCol).value = header
            nCol += 1
            if nCol > 6:
                break
        sheet.cell(row = nRow, column = 7).value = createRichText(requirement,targetstatements)

        nRow += 1

    (root,ext) = os.path.splitext(filename)
    newfile = root+"_TCA.xlsx"
    wb.save(newfile)
    print ("Created %s" % newfile )

def determineRunColor(run):
    if run.font.color.type is None:
        ret = (run.style.font.color.type,run.style.font.color.rgb)
    else:
        ret =  (run.font.color.type,run.font.color.rgb)
    return ret

def isTextCR(text,paragraph):
    ret = (False)
    for run in paragraph.runs:
        if not (run.font.strike or run.style.font.strike):
            if  run.text in text or text in run.text:
                #run is part of matched text, so lets check if its red
                #print(run.style.font.color.rgb)
                (type,rgb) = determineRunColor(run)
                if type == MSO_COLOR_TYPE.RGB:
                    if rgb[0] > 0:
                        ret = (True)
                        break
    return ret


def processNonHeader(paragraph,outlineLevel,headers,delegate):
    text = getAcceptedText(paragraph)
    if not text: 
        return
    
    targetstatements = ("shall","may")

    sentences = nltk.sent_tokenize(text)

    for group in sentences:
        for targetword in targetstatements:
            regex="\W" + targetword + "\W"    
            match= re.search(regex,group)
            if match:
                #if isTextCR(group,paragraph):
                delegate(outlineLevel,headers,group)
                break


def parseDocument(doc, delegate,addelegate):

    ht = HeadingTracking()

    for element in doc.iter_inner_content():
        if isinstance(element, docx.text.paragraph.Paragraph):
            processParagraph(element,ht,delegate,addelegate)
        elif isinstance(element, docx.table.Table):
            if not ht.ignoreableSection:
                processTable(element,ht,addelegate)
        else :
            print (element)

def processParagraph(paragraph,ht,nonheadingdelegate,addelegate):
    text = getAcceptedText(paragraph)
    if "Heading" in paragraph.style.name:
        if "Heading RevTable" in paragraph.style.name:
            ht.ignoreableSection = True
            pass
        if "Apx Heading" in paragraph.style.name:
            ht.ignoreableSection = True
            pass
        elif "Heading 8" in paragraph.style.name or "Heading 9" in paragraph.style.name:
            pass
        else:
            ht.ignoreableSection = False
            ht.pushHeader(paragraph.style,text)
    else:
        if not ht.ignoreableSection:
            nonheadingdelegate(paragraph,ht.getOutlineLevel(), ht.getHeaders(), addelegate)
#            print( "%d %s" % (style2level(paragraph.style), text))




def style2level(style):
    ret = 0
    match= re.search("\d+",style.name)
    if match:
        ret = int(match.group())
    return ret

def mergeText(text, newtext):
    if text:
        text += "-" + newtext
    else:
        text = newtext
    return text 

def processTable(table,ht,delegate):
    if len(table.rows) <= 1 or len(table.columns) <= 1:
        return
    
    magickeywords = ("\\bRequirement\\b" ,  "\\bSupport\\b" , "\\bStatus\\b")
    checkableColumn = {}

    bFirstRow = True
    for (r, row) in enumerate(table.rows):
        if bFirstRow:
            for (c, cell) in enumerate(row.cells):
                for keyword in magickeywords:
                    match = re.search(keyword,cell.text)
                    if match:
                        #print (cell.text)
                        checkableColumn[c]= c
            bFirstRow = False
            if len(checkableColumn) ==  0:
                break
        else:
            text = ""
            value = ""
            for (c, cell) in enumerate(row.cells):
                if c in checkableColumn:
                    value = row.cells[c].text
                else:
                    text = mergeText(text,row.cells[c].text)
            text = mergeText(text, value)
            delegate(ht.getOutlineLevel(), ht.getHeaders(),text)
            #for column in checkableColumn:
            #    print (row.cells[column].text)

    if len(checkableColumn) > 1:
        print (bcolors.FAIL + "Table in %s has more than one checkable column" % ht.getOutlineLevel() + bcolors.ENDC )
