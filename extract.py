##################################################
## extract.py
##################################################
## Author: gene@cloud2gnd.com
## Copyright: Copyright 2025
## Version: 1.0
##################################################

import docx.table
import docx.text
import docx.text.paragraph
import docx.shared

from bcolors import bcolors
from wordhelper import getAcceptedText
from excelhelper import setHeaderCell

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT,WD_SECTION
from docx.enum.dml import MSO_COLOR_TYPE
from docx.shared import Twips

import openpyxl
from openpyxl.cell.text import InlineFont
from openpyxl.formatting.rule import CellIsRule, FormulaRule
from openpyxl.cell.rich_text import TextBlock, CellRichText
import docx
import os
import re
import nltk
import itertools
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

numberToString = "Zero", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight"
#import ssl
#try:
#    _create_unverified_https_context = ssl._create_unverified_context
#except AttributeError:
#    pass
#else:
#    ssl._create_default_https_context = _create_unverified_https_context

class HeadingTracking:

    def __init__(self) -> None:
        self.lastHeaderLevel = 0
        self.outlinelevel={}
        self.headers={}
        self.ignoreableSection = False
        self.testcase = False  # True if we are in a test case section, so we ignore all headers
        self.testprocedure = False  # Added for Test Procedure tracking
        pass

    def __incrementOrSet__(self,level):
        if level in self.outlinelevel:
            self.outlinelevel[level] += 1
        else:
            self.outlinelevel[level] = 1

    def pushHeader(self,style,text):
        level=style2level(style)
        if level > self.lastHeaderLevel:
            #indent
            self.__incrementOrSet__(level)
        elif level < self.lastHeaderLevel:
            #undent
            #reset all sub to zero
            self.__incrementOrSet__(level)
            for x in range(level+1,8):
                foo = self.outlinelevel.pop(x,"foo")
                bar = self.headers.pop(x,"bar")
                #self.outlinelevel[level] = 0
        else:
            #level is the same
            self.__incrementOrSet__(level)

        self.headers[level] = text
        self.lastHeaderLevel = level

    def getHeaders(self):
        return self.headers.copy()
    
    def getOutlineLevel(self):
        text = ""
        bFirst = True
        for x in range (1,7):
            if x in self.outlinelevel:
                text = text + "%s%s" % ("." if x > 1 else "" , self.outlinelevel[x])
        return text

################################################################################################################################################

def extractTS(filename):
    print("extractTS:", filename)
    if  not os.path.exists(filename):
        print(bcolors.FAIL + "File '%s' does not exist" % filename + bcolors.ENDC)
        return
    doc = Document(filename)
    testText = []
    parseTestDocument(doc, nulldelegate, lambda x, y: testText.append((x, y)) )
    print("done")
    print(testText)
    if filename.find(".tex") != -1:
        outputFile = filename.le
    outputTeX(filename, testText)

def isTextCR(text,paragraph):
    ret = (False)
    for run in paragraph.runs:
        if not (run.font.strike or run.style.font.strike):
            if  run.text in text or text in run.text:
                #run is part of matched text, so lets check if its red
                #print(run.style.font.color.rgb)
                (type,rgb) = determineRunColor(run)
                if type == MSO_COLOR_TYPE.RGB:
                    if rgb[0] > 0:
                        ret = (True)
                        break
    return ret

def nulldelegate():
    return

def processNonHeader(paragraph,outlineLevel,headers,delegate):
    text = getAcceptedText(paragraph)
    if not text: 
        return
    
    targetstatements = ("shall","may")

    sentences = nltk.sent_tokenize(text)

    for group in sentences:
        for targetword in targetstatements:
            regex="\W" + targetword + "\W"    
            match= re.search(regex,group)
            if match:
                #if isTextCR(group,paragraph):
                delegate(outlineLevel,headers,group)
                break


def parseTestDocument(doc, delegate,adddelegate):

    ht = HeadingTracking()

    for element in doc.iter_inner_content():
        if isinstance(element, docx.text.paragraph.Paragraph):
            processTestParagraph(element,ht,delegate,adddelegate)
        elif isinstance(element, docx.table.Table):
            print("table")
            if not ht.ignoreableSection:
                print("not ht.ignorableSection")
                rows = convertTestTable(element, ht)
                if rows:
                    processTestTable(rows, adddelegate)
                
            else:
                print("ht.ignorableSection")
        else :
            print (element)

def processTestParagraph(paragraph,ht,nonheadingdelegate,adddelegate):
    text = getAcceptedText(paragraph)
    print("Style:", paragraph.style.name, "Text: ", text)
    adddelegate(paragraph.style.name, text)
    
    return
    if "Heading" in paragraph.style.name:
        if "Heading RevTable" in paragraph.style.name:
            ht.ignoreableSection = True
            pass
        if "Apx Heading" in paragraph.style.name:
            ht.ignoreableSection = True
            pass
        elif "Heading 8" in paragraph.style.name or "Heading 9" in paragraph.style.name:
            pass
        else:
            ht.ignoreableSection = False
            ht.pushHeader(paragraph.style,text)
    else:
        if not ht.ignoreableSection:
            nonheadingdelegate(paragraph,ht.getOutlineLevel(), ht.getHeaders(), addelegate)
#            print( "%d %s" % (style2level(paragraph.style), text))




def style2level(style):
    ret = 0
    match= re.search("\d+",style.name)
    if match:
        ret = int(match.group())
    return ret

def mergeText(text, newtext):
    if text:
        text += "-" + newtext
    else:
        text = newtext
    return text 

def isShaded(cell):
    retVal = False
    pattern = re.compile('w:fill=\"(\S*)\"')
    match = pattern.search(cell._tc.xml)
    if match:
        print("fill:", match.group(1), " type:", type(match.group(1)))
        if match.group(1) != "FFFFFF":
            retVal = True
    return retVal

def emToCM(ems):
    if(ems != None):
        return  (ems / 914400.0) * 2.54 #914400 em/in
    else:
        return 0

def indentSpaces(num):
    spaces = ' ' * num * 4
    return spaces

def processTestTable(rows, adddelegate):
    outText = ""
    firstRow = True

    firstRowHeader = "\\StartTable{|}"
    headerRowStart = "\\StartHeaderRow\n"
    embeddedRowStart = "\\StartEmbeddedHeaderRow\n"
    nonHeaderRowStart = "\\StartTableRow\n"
    headerRowEnd = "\\EndHeaderRow\n"
    embeddedRowEnd = "\\EndEmbeddedHeaderRow\n"
    processingFirstRow = True
    for row in rows:
        rowOut = ""
        isHeaderRow = row[0] == 1
        for col in itertools.islice(row, 1, None):
            if row[0] == 1:
                if isHeaderRow:
                    if processingFirstRow:
                        firstRowHeader = firstRowHeader[:len(firstRowHeader) - 1] #removing the "}" at the end. In order to not have to add an ending }
                                                                                  #after we're done processing, we add it every time and when we add a new 
                                                                                  #column width, we need to remove the last "}"
                        val = col[1]
                        firstRowHeader += " L{" + str(val) + " cm}|}"   #adding a } at the end so taht we don't have to proactively add it when we're done
                                                                        #processing the end of the first row.
                    if len(rowOut) == 0:
                        rowOut += headerRowStart
                        headerRowStart = embeddedRowStart
                    isHeaderRow = True
                    if(processingFirstRow):
                        rowOut += indentSpaces(1) + "\\HeaderCell{" + col[0] + "} & \n"
                    else:
                        rowOut += indentSpaces(1) + "\\HeaderCell{" + col + "} & \n"
                else:
                    isHeaderRow = False
                    if len(rowOut) == 0:
                        rowOut += nonHeaderRowStart
                        rowOut += indentSpaces(1) + "\\TableCell{"+col[1]+"}"
                    else:
                        rowOut += " &\n" + indentSpaces(1) + "\\TableCell{"+col[1]+"}"
                continue
            else:
                if len(rowOut) == 0:
                    rowOut += nonHeaderRowStart
                    rowOut += indentSpaces(1) + "\\TableCell{"+col+"}"
                else:
                    rowOut += " &\n"+ indentSpaces(1) + "\\TableCell{"+col+"}"
        if isHeaderRow:
            #for the header row, we want to do the opposite of what we did with firstRowHeader.  Here we want to remove the "&" at the end of the row
            #only for the last one.  This might not be the best way but it works for now.
            rowOut = rowOut[:len(rowOut) - 3]
            rowOut += "\n" + headerRowEnd
            headerRowEnd = embeddedRowEnd

        else:
            rowOut += "\n\\EndTableRow\n"
        if processingFirstRow:
            firstRowHeader += "\n"
            outText += firstRowHeader
            processingFirstRow = False
        processingFirstRow = False
        outText += rowOut
    outText += "\\TableCaption{caption}{tab:}\n\\EndTable"
    adddelegate("Table", outText)
    print(outText)
    return

def convertTestTable(table,ht):#,adddelegate):
    if len(table.rows) <= 1 or len(table.columns) <= 1:
        return
    
    magickeywords = ("\\bRequirement\\b" ,  "\\bSupport\\b" , "\\bStatus\\b")
    checkableColumn = {}

    bFirstRow = True
    rows = []
    numColumns = len(table.columns)
    processedShaded = False
    print("table style:", table.style)
    for (r, row) in enumerate(table.rows):
        print("r:", r)
        print("num rows:", len(table.rows))
        tableRow = []
        processedShadedForRow = False
        if bFirstRow:
            print("firstRow")
            for (c, cell) in enumerate(row.cells):
                print("c:", c, " cell:", cell.text, "width:", cell.width, " width type:", type(cell.width))
                if processedShadedForRow == False:
                    print("is summary row")
                    if isShaded(cell):
                        tableRow.append(1)
                    else:
                        tableRow.append(0)
                    processedShadedForRow = True
                print("width type:", type(cell.width))
                print("width type:", type(type(cell.width)))
                print("twips type:", type(Twips))
                #print("width in cm", cell.width.cm)
                if(cell.width == None):
                    cellWidth = 16.0 / numColumns #since the table row doesn't include column widths, just use 16 cm (approx 6.25") / num columns
                    tableRow.append([cell.text, cellWidth]) 
                else:
                    tableRow.append([cell.text, emToCM(cell.width)])
            bFirstRow = False
            print(tableRow)
            rows.append(tableRow)
            #if len(checkableColumn) ==  0:
                #break
        else:
            value = ""
            tableRow = []
            processedShadedForRow = False
            for (c, cell) in enumerate(row.cells):
                print("c:", c, " cell:", cell.text)
                print("v:", row.cells[c].text)
                if processedShadedForRow == False:
                    print("is summary row")
                    if isShaded(cell):
                        tableRow.append(1)
                    else:
                        tableRow.append(0)
                    processedShadedForRow = True
                tableRow.append(cell.text)
            rows.append(tableRow)
            print(tableRow)
            print(rows)
    #adddelegate("Table", rows)
    return rows
    #if len(checkableColumn) > 1:
    #    print (bcolors.FAIL + "Table in %s has more than one checkable column" % ht.getOutlineLevel() + bcolors.ENDC )

def headingStyle(processingState, styleType):
#   Heading Style is just '\Heading'+number as a word.  ex. 'Heading 2' = '\HeadingTwo'
    print("Heading Style: ", styleType)
    outText = ""
    if processingState["inTestSteps"]:
        outText += "\\EndTestSteps\n"
        processingState["inTestSteps"] = False

    outText = clearProcessingState(processingState, outText)
    processingState["lastSectionWasHeading"] = True

    match = re.match(r"\s*heading.(.*)\s*", styleType.lower())
    if match:
        if processingState["inTest"]:
            print("exiting test steps")
            outText += "\\end{itemize}\n"
            processingState["inTest"] = False
        if int(match.group(1)) == 8:
            print("in test steps")
            lastHeadingNum = processingState["lastHeadingNumber"]
            outText += "\Heading"+numberToString[lastHeadingNum + 1]+"{%s}"
            processingState["inTest"] = True
        else:
            processingState["lastHeadingNumber"] = int(match.group(1))
            outText += "\Heading"+numberToString[int(match.group(1))]+"{%s}"
    else:
        outText += "Heading # not found (%s)"
    return (outText)


def listStyle(processingState, styleType):
    print("List Style: '", styleType, "' firstList:", processingState["firstList"], "inTestSteps:", processingState["inTestSteps"])
    outText = ""
    match = re.match(r"\s*list\snumber\s(.*)\s*", styleType.lower().strip())
    if match:
        print("match")
        if int(match.group(1)) == 2:
            print("in test steps")
            if processingState["inTestSteps"] == False:
                outText += "\\StartTestSteps\n"
                processingState["inTestSteps"] = True;
            outText += "\\TestStepItem{%s}"
            return outText
        else:
            print("no regex match")
    else:
        print("no match")
    if processingState["inTestSteps"]:
        outText += "\\EndTestSteps\n"
        processingState["inTestSteps"] = False
    if processingState["firstList"]:
        print("exit test steps")
        processingState["firstList"] = False
        outText += "\\begin{itemize}\n\\item{%s}"
    else:
        outText += "\\item{%s}"
    return outText


def normalStyle(processingState, styleType):
    print("Normal")
    outText = ""
    outText = clearProcessingState(processingState, outText)
    outText += "%s"
    return (outText)

def testCaseStyle(processingState, styleType):
# There are 3 'Test Case" styles.  'Test Case Heading' is a \TestItem. 'Test Case Verdict' is '\TestContent{\uline{text}}' (This could change in the future to have a dedicated style).  'Test Case Body' is '\TestContent'
    print("Test Case: ", styleType)
    outText = ""
    if processingState["inTestSteps"]:
        outText += "\\EndTestSteps"
        processingState["inTestSteps"] = False
    wasLastSectionHeading = processingState["lastSectionWasHeading"]
    outText = clearProcessingState(processingState, outText)
    if(styleType.lower().find("heading")) != -1:
        if wasLastSectionHeading:
            outText += "\\begin{itemize}\n"
            processingState["inTest"] = True  #need to set this here separately for test driven tests.  The single tests have a different heading number (8)
        outText += "\\TestItem{%s}"
    elif(styleType.lower().find("verdict")) != -1:
        outText += "\\TestContent{\\uline{%s}}"
    elif(styleType.lower().find("body")) != -1:
        outText += "\\TestContent{%s}"
    else:
        outText += "NOT FOUND(%s)"
    return outText

def referenceStyle(processingState, styleType):
    print("Reference")
    if processingState["firstReference"]:
        processingState["firstReference"] = False
    return ("\\DefineReference{}{%s}")

def bodyStyle(processingState, styleType):
    print("body")
    outText = ""
    outText = clearProcessingState(processingState, outText)
    outText += "%s"
    return (outText)

def captionStyle(processingState, styleType):
    print("body")
    outText = ""
    outText = clearProcessingState(processingState, outText)
    outText += "caption style(%s)"
    return (outText)

def disclaimerStyle(processingState, styleType):
    print("disclaimer")
    outText = ""
    outText = clearProcessingState(processingState, outText)
    outText += "\\ShortDisclaimer{2025}"
    return (outText)

def tableStyle(processingState, styleType):
    print("table")
    return ("%s")

def convertWordStyleToLaTeXStyle(processingState, styleType, styleDelegateList):
    print ("styleDelegateList:", styleDelegateList)
    for [type, delegate] in styleDelegateList:
        if styleType.lower().find(type) != -1:
            return delegate(processingState, styleType)
    print("not found:", styleType)
    return ("style not found(%s)".format(styleType))
#    if styleType.lower().find("Test Case Body") != -1:
#        print("test case body")
#    else if styleType.lower().find

def clearProcessingState(processingState, outText):
    if processingState["firstList"] == False:
        print("add end itemize")
        outText += "\\end{itemize}\n"
        print(outText)
    processingState["firstList"] = True
    processingState["firstReference"] = True
    processingState["lastSectionWasHeading"] = False
    return outText

def outputTeX(filename, testText):
    styleDelegateList = [["test case", testCaseStyle], ["list", listStyle], ["heading", headingStyle], ["normal", normalStyle], ["body",bodyStyle], ["caption", captionStyle], ["reference", referenceStyle], ["disclaimer", disclaimerStyle], ["table", tableStyle], ["lastHeadingNumber", 0]]
    processingState = {"firstList": True, "firstReference": True, "inTest": False, "inTestSteps": False, "lastSectionWasHeading": False}
    print("firstList:", processingState["firstList"])
    print("firstReference:", processingState["firstReference"])

    print("orig file:", filename)
    match = re.match(r"(.+?)(\.[^.]*$|$)", filename) #(r"\s*(.*)(docx)\w", filename)
    if match:
        outputfile = match.group(1) + ".tex"
    else:
        outputfile = "unknown.tex"

    fp = open(outputfile, "w")
    fp.write("\\documentclass{bluetooth.test}\n")
    fp.write("\\SetSpecificationName{Electronic Shelf Label Service (ESLS)}\n")
    fp.write("\\SetBluetoothDocumentType{Test Suite}\n")
    fp.write("\\SetRevision{d09r01}\n")
    fp.write("\\SetGroup {Electronic Shelf Label Working Group}\n")
    fp.write("\\SetFeedback {esl-main@bluetooth.org}%\href{mailto:esl-main@bluetooth.org}{esl-main@bluetooth.org}}\n")
    fp.write("\\BluetoothDraftSpec\n")
    fp.write("\\SetAbstract {This service allows electronic shelf labels (ESLs) to be controlled and updated using Bluetooth wireless technology.}\n")
    fp.write("\\begin{document}\n")
    fp.write("\\FrontCover\n")
    fp.write("\\clearpage\n")
    for (styleType, text) in testText:
        #print(convertWordStyleToLaTeXStyle(styleType, styleDelegateList) % styleType)
        outText = convertWordStyleToLaTeXStyle(processingState, styleType, styleDelegateList)
        if outText.find("%s") != -1:
            print(outText % text)
            fp.write(outText % text + "\n")
        else:
            print(outText)
            fp.write(outText)
    fp.write("\end{document}\n")
    fp.close()
    print ("outputfile: ", outputfile)
    
#        print("\\HeadingTwo{%r}" % styleType)

