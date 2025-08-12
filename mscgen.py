
import re

# mscgen.py now contains processMyDocx; mydocx.py is obsolete and can be deleted.
import os
import re
from docx import Document
import docx
import nltk
from enum import Enum

from bcolors import bcolors
from tca import HeadingTracking
from wordhelper import getAcceptedText
from nltk.parse import RecursiveDescentParser

class ActionType(Enum):
    UNKNOWN = 0
    LTR = 1
    RTL = 2
    BOTH = 3
    NONE = 4
    BOX = 5
    BEGINRECT = 6
    ENDRECT = 7
    # Add more as needed

class NodeRole(Enum):
    UNKNOWN = 0
    VERBFOUND = 1
    TARGET = 2
    ACTION = 3
    ITEM = 4
    EMPTYVERB = 5
    # Add more as needed

class MyNode:
    def __init__(self, source="",target=""):
        self.source = source
        self.target = target
        self.action = ActionType.NONE
        self.subject = ""
        self.indentContext = ""
        self.parmeters = []


class TestStepsTracker:
   
    
    def __init__(self, TCID ="") -> None:
        self.teststeps = []  # Changed from string to list
        self.currentTCID=TCID
        self.indentContext= ""
        pass

    def reset(self, TCID=""):
        self.teststeps = []
        self.currentTCID = TCID
        self.indentContext = ""

    def __incrementOrSet__(self,level):
        if level in self.outlinelevel:
            self.outlinelevel[level] += 1
        else:
            self.outlinelevel[level] = 1

    def pushTestStep(self,style,text):
        # Only add text if style is List Number 2 or Caption
        if style == "List Number 2":
            self.teststeps.append({"type": "li", "step": text})
        elif style == "Caption":
            pass
        else:
            self.teststeps.append({"type": "", "step": text})
        pass

    def fix_pos_tags(self, pos_tags):
        """
        Fix XX being tagged as VB, and retag ('set', 'NN') to ('set', 'VB') when followed by ('to', 'TO').
        """
        # fix XX being tagged as VB
        pos_tags = [(word, 'NN') if word == 'XX' else (word, tag) for word, tag in pos_tags]
        # If ('set', 'NN') is followed by ('to', 'TO'), change 'set' to ('set', 'VB')
        new_pos_tags = []
        i = 0
        while i < len(pos_tags):
            word, tag = pos_tags[i]
            if (
                word == 'set' and tag == 'NN' and
                i + 1 < len(pos_tags) and pos_tags[i + 1][0] == 'to' and pos_tags[i + 1][1] == 'TO'
            ):
                new_pos_tags.append(('set', 'VB'))
            else:
                new_pos_tags.append((word, tag))
            i += 1
        return new_pos_tags


    def getSteps(self):
        processed = []
        for item in self.teststeps:
            if item.get("type") != "li":
                continue

            step_text = item.get("step", "")
            step_text = replace_bracketed_indices(step_text)  # Replace bracketed indices
            step_text = replace_comparison_symbols(step_text)  # Replace comparison symbols

            # Use nltk to tokenize and POS tag the step
            tokens = nltk.word_tokenize(step_text)


            pos_tags = nltk.pos_tag(tokens)
            pos_tags = self.fix_pos_tags(pos_tags)

            # Define grammar for chunking
            grammar = r"""
                SECTION: {<NNP|NN|NNS><CD>}                                      # Section 3.3.3.1 as a single chunk if CD is one token
                SECTION: {<NNP|NN|NNS><CD>(<\.|:|,><CD>)+}                      # Section 3.3.3.1 as a sequence (legacy, if split)
                TABLE: {<IN><VBN><IN><JJ><CD>}                                      # Table
                NP: {<DT>?<JJ.*>*<NN.*|NNS|NNP|NNPS|FW>+<BRACKET\w+>?}           # Noun phrase
                NP: {<NN.*|NNS|NNP|NNPS|FW>+<\[\w*\]>}                         # Noun phrase with bracketed index (e.g., HCI_LE_Set_CIG_Parameters [v2])
                NP: {<NN.*|NNS|NNP|NNPS|FW>+<BRACKET\w+>}                        # Noun phrase with trailing number (e.g., Config_ID 1)
                VP: {<VB.*><TO><NN.*|NNS|NNP|NNPS|FW|CD|JJ.*|RB.*>+}              # Verb phrase with 'set to XX'
                VP: {<VB.*><TO><NP|CD>+}  
                CONJCONJ: {<CC><VB.*>}
                CONJ: {<CC>}                                                    # Conjunction
                NUM: {<CD>}                                                     # Numbers
                EVENT: {<NNP>+}                                                 # Named events (e.g., HCI_Command_Complete)
                CLAUSE: {<NP><VP>}                                              # Clause (NP + VP)
                COORD_NP: {<NP>(<,>?<CC><NP>)+}                                 # Coordinated NPs
                COORD_VP: {<VP>(<,>?<CC><VP>)+}  
            """
            chunk_parser = nltk.RegexpParser(grammar)
            tree = chunk_parser.parse(pos_tags)
#                CLAUSE: {<NP><VP>}                                              # Clause (e.g., NP followed by VP)
#                VP: {<VB.*><NP|PP|CLAUSE>+}                                     # Verb phrase
 #               NP: {<NN.*|NNS|NNP|NNPS|FW>+<\\[><CD><\\]>}                     # Noun phrase with bracketed number (e.g., Coded_Rates_C_To_P [ 0 ])
 # PP: {<IN><NP>}                                                  # Prepositional phrase (e.g., to the Lower Tester)
  #              PP: {<IN><NP><,>?<CC>?<NP>}                                     # PP with coordination (e.g., with X and Y)
   #             PP: {<IN><NP><,>?<NP>} 
            """"
            rules = nltk.data.load('brown', 'text')
            grammar = nltk.CFG.fromstring(rules)
            #grammar = nltk.data.load('grammars/large_grammars/atis.cfg', 'text')
            rd = RecursiveDescentParser(grammar)
            tree = rd.parse(tokens)
            """

            things = self.iterate_all_nodes(tree)  # Process all nodes in the tree
            processed.extend(things)

        return processed
    
    def parseVerb(self, verb):
        """
        Parse a verb and return its action and item.
        This is a placeholder for actual parsing logic.
        """
        ret = ActionType.BOTH

        # Example parsing logic, can be replaced with actual logic
        if verb in ( "sends", "send", "generates"):
            ret =  ActionType.LTR
        elif verb in ("execute" , "receives" ,"returns" ):
            ret = ActionType.RTL
        elif verb == "perform" or verb == "forces" :
            ret = ActionType.BOX
        elif verb == "verifies":
            ret = ActionType.BOX
        else:
            print(bcolors.FAIL + "Unknown verb: %s" % verb + bcolors.ENDC)
            ret= ActionType.UNKNOWN

        return ret
        
    def iterate_all_nodes(self, tree):
        """
        Given an NLTK tree, call preterminal_action(node) for preterminals and leaf_action(node) for leaves.
        A preterminal is a node whose children are all leaves.
        """
        things= []
        foo = MyNode()
        ctx = NodeRole.UNKNOWN
        previousNoun = ""
        previousVerb = ""
        isFirstStep = True

        for child in tree:
            if isinstance(child, nltk.Tree):
                treelabel = child.label()
                if treelabel == "NP":
                    previousNoun = tree2String(child)

                    if ctx == NodeRole.TARGET:
                        foo.target = previousNoun
                        ctx = NodeRole.UNKNOWN
                    elif ctx == NodeRole.EMPTYVERB:
                        foo.subject = previousNoun
                        ctx = NodeRole.UNKNOWN
                    elif ctx == NodeRole.VERBFOUND:
                        if isTestRole(previousNoun):
                            foo.target = previousNoun
                            foo.subject = previousVerb
                        else:
                            foo.subject = previousNoun
                        previousNoun =""
                        ctx = NodeRole.UNKNOWN
                    else:
                        #print("unknown context for '%s'" % previousNoun)
                        pass
                elif treelabel == "NUM":
                    if isFirstStep:
                        self.indentContext = tree2String(child)
                    else:
                        print(bcolors.WARNING + "Unknown Number : %s (%s)" % (treelabel,tree2String(child)) + bcolors.ENDC)
                        pass
                elif treelabel == "CONJ":
                    print(bcolors.WARNING + "Ignore conjuction : %s (%s)" % (treelabel,tree2String(child)) + bcolors.ENDC)
                    pass
                elif treelabel == "PP":
                    print(bcolors.WARNING + "Ignore PP : %s (%s)" % (treelabel,tree2String(child)) + bcolors.ENDC)
                    pass
                elif treelabel == "CLAUSE":
                    np, nn = extract_np_nn_from_tree(child)
                    if np and nn:
                        foo.parmeters.append((np, nn))  # Store parameters
                    else:
                        print(bcolors.WARNING + "CLAUSE without NP or NN: %s" % tree2String(child) + bcolors.ENDC)
                elif treelabel =="CONJCONJ":
                    things.append(foo)
                    foo = MyNode(foo.source,foo.target)
                    foo.action = ActionType.LTR
                    ctx = NodeRole.VERBFOUND
                elif treelabel == "VP":
                    foo.subject = tree2String(child)
                elif treelabel == "SECTION":
                    pass  # Ignore SECTION nodes
                elif treelabel == "TABLE":
                    pass  # Ignore SECTION nodes
                else:
                    print(bcolors.FAIL + "Unknown tree label: %s" % treelabel + bcolors.ENDC)
                    pass
            else :
                label = child[1]
                if label == "VBZ" or label == "VB" or label == "VBP" or label == "VBD" or label == "VBG":
                   # if label == "VBZ":
                    #    print(bcolors.WARNING + "found VBZ: %s" % child[0] + bcolors.ENDC)

                    verb = child[0].lower()
                    if verb == "set" or verb == "does" or verb == "steps" or verb == "depending":
                        print(bcolors.WARNING + "ignoring %s" % verb + bcolors.ENDC)
                        pass
                    else:
                        previousVerb = verb
                        action = self.parseVerb(verb)
                        print(f"Action for verb '{child[1]}': {action}")
                        if action != ActionType.UNKNOWN:
                            foo.action = action
                    
                            ctx = NodeRole.VERBFOUND
                            if len(previousNoun) > 0:
                                foo.source = previousNoun
                            else:
                                ctx= NodeRole.EMPTYVERB
                                print(bcolors.WARNING + "Empty noun found in text: %s" % child[0] + bcolors.ENDC)
                        else:
                            print(bcolors.FAIL + "Unknown verb: %s" % verb + bcolors.ENDC)
                            ctx = NodeRole.UNKNOWN
                elif label == "TO":
                    #foo.subject = previousNoun
                    ctx = NodeRole.TARGET
                elif label == "CC" :
                    foo.indentContext = self.indentContext
                    things.append(foo)
                    foo = MyNode(foo.source,foo.target)
                    ctx = NodeRole.UNKNOWN
                    previousNoun = ""
                elif label == "IN":
                    pass # Ignore prepositions
                elif label in (".",",", ":", ";", "!", "?" , "(",")"):
                    pass # Ignore punctuation
                elif label == "CD":
                    pass # Ignore cardinal numbers
                elif label == "MD":
                    pass # Ignore modals
                elif label == "RB":
                    pass # Ignore modals
                else:
                     print (bcolors.FAIL + "Unknown label (%s)" % child[1] + bcolors.ENDC )
                     pass
            isFirstStep = False

        foo.indentContext = self.indentContext
        things.append(foo)  # Add the last node
        self.indentContext = ""  # Reset indentContext after processing
        return things

  
    



    def getOutlineLevel(self):
        text = ""
        bFirst = True
        for x in range (1,7):
            if x in self.outlinelevel:
                text = text + "%s%s" % ("." if x > 1 else "" , self.outlinelevel[x])
        return text

    def setCurrentTCID(self,text):
       if len(self.currentTCID) == 0:
            self.currentTCID= text 

    def getCurrentTCID(self):
        return self.currentTCID.strip() if self.currentTCID else "[TCID]"

##################################################

def generateLanes(things, f):
    ut = False
    iut = False
    lt = False

    for thing in things:
        src = thing.source.lower()
        tgt = thing.target.lower()
        if src == "upper tester" or src == "ut" or tgt == "upper tester" or tgt == "ut":
            ut = True
        if src == "iut" or tgt == "iut":
            iut = True
        if src == "lower tester" or src == "lt" or tgt == "lower tester" or tgt == "lt":
            lt = True

    lanes=[]
    offset = 2
    if ut:
        f.write("vertical UT { label = \"Upper Tester\", x = %scm }\n" % offset)
        offset += 6
        lanes.append("UT")
    if iut:
        f.write("vertical IUT { label = \"IUT\", x = %scm }\n" % offset)
        offset += 6
        lanes.append("IUT")
    if lt:
        f.write("vertical LT { label = \"Lower Tester\", x = %scm }\n" % offset)
        offset += 6
        lanes.append("LT")

    f.write("\n")

    firstlane = lanes[0] if lanes else None
    lastlane = lanes[-1] if lanes else None
    return (firstlane, lastlane)

# Utility function to mangle subject strings for MSC output
def mangleSubject(subject):
    """
    Modify subject string for MSC output.
    - Change 'successful NP event' to 'NP event (success)'
    - Extend with more rules as needed.
    """
    import re
    # Rule 1: 'successful NP event' -> 'NP event (success)'
    m = re.match(r'successful\s+(.+?\s*event)', subject, re.IGNORECASE)
    if m:
        return f"{m.group(1).strip()} (success)"
    # Rule 2: 'successful NP' -> 'NP (success)' (not just event)
    m2 = re.match(r'successful\s+(.+)', subject, re.IGNORECASE)
    if m2:
        return f"{m2.group(1).strip()} (success)"
    # If not matched, return original
    return subject

def generateMSC(reqlist, path):
    """
    Placeholder function for generateMSC.
    Implement the logic to generate MSC from reqlist and path.
    """

    # Ensure /msc/ subdirectory exists
    msc_path = os.path.join(path, "msc")
    os.makedirs(msc_path, exist_ok=True)
    for things in reqlist:
        # Check that things is a non-empty list of MyNode objects
        if isinstance(things[2], list):
            tcid = str(things[1]).strip()
            if tcid == "[TCID]":
                tcid = str(things[0]).strip()
            # Replace /, [, ] with _
            for ch in ['/', '[', ']']:
                tcid = tcid.replace(ch, '_')
            filename = os.path.join(msc_path, f"{tcid}.txt")
            with open(filename, "w", encoding="utf-8") as f:
                f.write(f"#\n#{tcid}\n#\n\n")
                f.write("include \"btsig.style\"; \nbegin msc; \n\n")
                # Write lane output using generateLanes
                (firstlane,lastlane) = generateLanes(things[2], f)
                alt_state = False
                previous_alt_state = False
                for thing in things[2]:
                    subject = reverse_bracketed_indices(thing.subject) if thing.subject else ""
                    # Check indentContext for patterns like XA.1, 6A.1, 6A.2, etc.
                    if thing.indentContext is not None and re.match(r'(?:\d+[A-Z])\.\d+', thing.indentContext.strip()):
                        alt_state = True
                        if previous_alt_state == False:
                            f.write("\n%s rect %s {label =\"%s\", align = left } \n" % ( source_string_to_lane(thing.source,firstlane ), source_string_to_lane(thing.target,lastlane), thing.indentContext.strip()))
                    elif previous_alt_state == True:
                        alt_state = False
                        f.write("end rect;\n\n")
                    else:
                        alt_state = False

                    #determine 'label' based on parameters or subject
                    if thing.parmeters is not None and len(thing.parmeters) > 0:
                        label = subject + reverse_bracketed_indices(buildParametersString(thing.parmeters))
                    else:
                        label = mangleSubject(subject)

                    if thing.source is None or len(thing.source) == 0 or thing.target is None or len(thing.target) == 0 or thing.action == ActionType.BOX:
                        action = "box"
                    else:
                        action = action_type_to_string(thing.action)

                    f.write(" %s %s %s {label =\"%s\" } \n" % (
                        source_string_to_lane(thing.source,firstlane ),
                        action,
                        source_string_to_lane(thing.target,lastlane),
                        label))
                        
                    previous_alt_state = alt_state
                if previous_alt_state == True:
                    f.write("end rect;\n")

                f.write("end msc;\n")

###################################################

def processMyDocx(path):
  print(f"Processing DOCX file: {path}")
  doc = Document(path)

  reqlist = []

  parseDocument(doc, processNonHeader, lambda x, y, z: reqlist.append((x, y, z)))

  dir_path = os.path.dirname(path)
  generateMSC(reqlist, dir_path)
  

def parseDocument(doc, delegate,addelegate):

    ht = HeadingTracking()
    ts= TestStepsTracker()

    for element in doc.iter_inner_content():
        if isinstance(element, docx.text.paragraph.Paragraph):
            processParagraph(element,ht,ts,delegate,addelegate)
#        elif isinstance(element, docx.table.Table):
#            if not ht.ignoreableSection:
#                processTable(element,ht,addelegate)
#        else :
#            print (element)


def processParagraph(paragraph,ht,ts,nonheadingdelegate,addelegate):
    text = getAcceptedText(paragraph)

    if "Heading" in paragraph.style.name:
        if "Heading RevTable" in paragraph.style.name:
            ht.ignoreableSection = True
            pass
        elif "Apx Heading" in paragraph.style.name:
            ht.ignoreableSection = True
            pass
        elif "Test Case Heading" in paragraph.style.name:
            if text == "Test Procedure":
                print ("### BEGIN OF TEST PROCEDURE ###")
                ht.testcase = True
                ht.testprocedure = True
                ts.reset(ts.getCurrentTCID())
                pass
            elif text == "Expected Outcome":
                if ht.testprocedure:
                    print ("### END OF TEST PROCEDURE ###")
                    print(ts.getCurrentTCID())
                    z= ts.getSteps()
                    addelegate(ht.getOutlineLevel(), ts.getCurrentTCID(), z)
                    ht.testcase = False
                    ts.reset()
                ht.testprocedure = False
            elif text == "Test Purpose":
                print("### TEST PURPOSE ###")
                pass
            elif text == "Reference":
                print("### REFERENCE ###")
                pass
            elif text == "Initial Condition":
                print("### INITIAL CONDITION ###")
                pass
            elif text == "Test Case Configuration":   
                print("### Test Case Configuration ###")
                pass
            elif text == "":   
                pass # ignore empty headings
            else:
                print(bcolors.FAIL + "Unknown Test Case Heading: %s" % text + bcolors.ENDC)
        elif "Heading 8" in paragraph.style.name:
            if ht.testcase:
                if ht.testprocedure:
                    print(ts.getCurrentTCID())
                    z= ts.getSteps()
                    addelegate(ht.getOutlineLevel(), ts.getCurrentTCID(), z)
                    ts.reset()
                    pass
                    #output stuff
            else:
                ht.testcase = True
            ts.setCurrentTCID(text)                
        elif "Heading 9" in paragraph.style.name:
            ts.setCurrentTCID(paragraph.style.name,text)
        else:
            ht.ignoreableSection = False
            ht.pushHeader(paragraph.style,text)
    elif "Test Case Verdict" in paragraph.style.name:
        pass
    else:
        if ht.testprocedure:
            ts.pushTestStep(paragraph.style.name, text)
        if not ht.ignoreableSection:
            nonheadingdelegate(paragraph,ht.getOutlineLevel(), ht.getHeaders(), addelegate)
            

def processNonHeader(paragraph,outlineLevel,headers,delegate):
    text = getAcceptedText(paragraph)
    if not text: 
        return
    
    targetstatements = ("shall","may")

    sentences = nltk.sent_tokenize(text)

    for group in sentences:
        for targetword in targetstatements:
            regex="\W" + targetword + "\W"    
            match= re.search(regex,group)
            if match:
                #if isTextCR(group,paragraph):
                delegate(outlineLevel,headers,group)
                break

def action_type_to_string(action_type):
    """
    Convert an ActionType enum value to a human-readable string.
    """
    if action_type == ActionType.LTR:
        return "->"
    elif action_type == ActionType.RTL:
        return "<-"
    elif action_type == ActionType.BOTH:
        return "<->"
    elif action_type == ActionType.NONE:
        return "<->"
    else:
        return "<->"

def source_string_to_lane(source,default=""):
    """
    Convert a source string to a lane identifier for MSC output.
    For example, 'upper tester' or 'UT' -> 'UT', 'iut' -> 'IUT', 'lower tester' or 'LT' -> 'LT'.
    Returns the lane string or the original string if no match.
    """
    s = source.strip().lower()
    if s in ("upper tester", "ut"):
        return "UT"
    elif s == "iut":
        return "IUT"
    elif s in ("lower tester", "lt"):
        return "LT"
    else:
        return default if default else s


def isTestRole(s):
    """
    Returns True if the string matches 'Lower Tester', 'Upper Tester', or 'IUT' (case-insensitive, allows abbreviations 'LT', 'UT', 'IUT').
    """
    s_norm = s.strip().lower()
    return s_norm in ("lower tester", "lt", "upper tester", "ut", "iut")

def replace_comparison_symbols(step_text):
    """
    Replace >, <, <=, >= in step_text with uppercase words: GT, LT, LE, GE.
    """
    # Replace >= and <= first to avoid partial replacement
    step_text = step_text.replace('>=', 'greater than or equal to')
    step_text = step_text.replace('<=', 'less than or equal to')
    step_text = step_text.replace('>', 'greater than')
    step_text = step_text.replace('<', 'less than')
    return step_text

def buildParametersString(parameters):
    """
    Given a list of parameters (typically tuples), format them for MSC output.
    Returns a string like ' [param1, param2, ...]' or '' if empty.
    """
    if not parameters:
        return ''
    # Flatten and stringify each parameter tuple
    param_strs = []
    for param in parameters:
        if isinstance(param, tuple):
            param_strs.append('='.join(str(p) for p in param))
        else:
            param_strs.append(str(param))
    return ' (' + '; '.join(param_strs) + ')'

def reverse_bracketed_indices(text):
    """
    Reverse the effect of replace_bracketed_indices: convert BRACKETxxx back to [xxx].
    """
    import re
    # Replace BRACKETxxx with [xxx], including cases like ACL_Connection_HandleBRACKET0 -> ACL_Connection_Handle[0]
    # Handles both BRACKETv2 and BRACKET0
    return re.sub(r'BRACKET(\w+)', r'[\1]', text)

def replace_bracketed_indices(text):
    """
    Replace all occurrences of [some string] (e.g., [0], [v2]) in the input text with a token like BRACKET0 or BRACKETv2.
    """
    def replacer(match):
        inner = match.group(1)
        return f"BRACKET{inner}"
    # Replace [something] with BRACKETsomething
    return re.sub(r'\[([^\]]+)\]', replacer, text)

def extract_np_nn_from_tree(tree):
    """
    Given an NLTK tree with the structure NP VB TO ... , return the NP and a list of all tokens after the TO in the VP.
    Returns (np_str, after_to_list) or (None, None) if the pattern does not match.
    """
    # The tree should be a parent node (e.g., S or CLAUSE) with children: NP, VP
    # We look for a pattern: [NP, VP], where VP is [VB, TO, ...]
    if not isinstance(tree, nltk.Tree) or len(tree) < 2:
        return (None, None)
    # Find NP and VP children
    np_subtree = None
    vp_subtree = None
    for child in tree:
        if isinstance(child, nltk.Tree):
            if child.label() == 'NP':
                np_subtree = child
            elif child.label() == 'VP':
                vp_subtree = child
    if np_subtree is None or vp_subtree is None:
        return (None, None)
    # VP should have children: VB, TO, ...
    vp_leaves = vp_subtree.leaves()
    if len(vp_leaves) < 3:
        return (None, None)
    # Check POS tags for VB and TO
    if vp_leaves[0][1].startswith('VB') and vp_leaves[1][1] == 'TO':
        np_str = tree2String(np_subtree)
        after_to_list = [w for (w, t) in vp_leaves[2:]]
        return (np_str, " ".join(after_to_list))
    return (None, None)

def tree2String( tree):
    """
    Convert an NLTK tree to a string representation, skipping any child that has a DT (determiner) POS tag.
    Removes leading and trailing spaces from the result.
    """
    if isinstance(tree, nltk.Tree):
        result = ' '.join(
            tree2String(child)
            for child in tree
            if not (
                (isinstance(child, tuple) and len(child) > 1 and child[1] == 'DT')
            )
        )
        return result.strip()
    else:
        # child is a tuple (word, tag)
        return tree[0].strip() if (len(tree) < 2 or tree[1] != 'DT') else ''