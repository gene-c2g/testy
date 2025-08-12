##################################################
## testy.py
##################################################
## Author: clenahan@cloud2gnd.com
## Copyright: Copyright 2024
## Version: 1.0
##################################################

from bcolors import bcolors
from wordhelper import Paras2Text
from excelhelper import setHeaderCell

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT,WD_SECTION
import openpyxl
from openpyxl.formatting.rule import CellIsRule, FormulaRule
import os

def processIOPTestPlan(filename):
    if  not os.path.exists(filename):
        print(bcolors.FAIL + "File '%s' does not exist" % filename + bcolors.ENDC)
        return
    
    doc = Document(filename)

    testlist = []

    parseTestPlan(doc, lambda w,x,y,z: testlist.append((w,x,y,z)))
    
    outputImportFile(filename,testlist)


def outputImportFile(filename,testlist):
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.sheet_view.zoomScale = 140
    setHeaderCell(sheet,1,"Category",15)
    setHeaderCell(sheet,2,"Test Case Identificier",20)
    setHeaderCell(sheet,3,"Test Case Name",30)
    setHeaderCell(sheet,4,"Feature",10)
    setHeaderCell(sheet,5,"Required",15)
    setHeaderCell(sheet,6,"Priority",15)
    setHeaderCell(sheet,7,"Pass Criteria",15)
    setHeaderCell(sheet,8,"Roles",15)

    nRow = 2
    for (rawtcid,iopid,passcnt,label) in testlist:
        tcid = rawtcid.split(" ")
        sheet.cell(row = nRow, column =1).value = iopid
        sheet.cell(row = nRow, column =2).value = tcid[0]
        sheet.cell(row = nRow, column =3).value = " ".join(tcid[1:]).replace("[","").replace("]","")
        sheet.cell(row = nRow, column =4).value = "XXX"
        sheet.cell(row = nRow, column =5).value = "Optional" if passcnt == 0 else "Mandatory"
        sheet.cell(row = nRow, column =6).value = 1
        sheet.cell(row = nRow, column =7).value = 1 if passcnt == 0 else passcnt
        sheet.cell(row = nRow, column =8).value = "IUT,PTS"
        nRow += 1

    (root,ext) = os.path.splitext(filename)
    newfile = root+"_IMPORT.xlsx"
    wb.save(newfile)
    print("Created %s" % newfile)

def parseTestPlan(doc, delegate):
    for table in doc.tables:
        if isTestTable(table):
            bFirst = True
            for r, rowiteration in enumerate(table.rows):
                if bFirst:
                    bFirst = False
                    continue
                if rowiteration.cells[0].grid_span == 1:
                    nextcell = 0
                    tcid = Paras2Text(rowiteration.cells[nextcell].paragraphs)
                    nextcell += rowiteration.cells[nextcell].grid_span
                    iopid = Paras2Text(rowiteration.cells[nextcell].paragraphs)
                    nextcell += rowiteration.cells[nextcell].grid_span
                    strPassCnt = Paras2Text(rowiteration.cells[nextcell].paragraphs)
                    if (strPassCnt.isdigit()):
                        passcnt = int(strPassCnt)
                    else:
                        passcnt = 0
                        print(bcolors.FAIL + "TCID %s (%s) has a blank pass count" % (tcid ,iopid)+ bcolors.ENDC)
                
                    nextcell += rowiteration.cells[nextcell].grid_span
                    label = Paras2Text(rowiteration.cells[nextcell].paragraphs)
                    nextcell += rowiteration.cells[nextcell].grid_span
                    delegate(tcid,iopid,passcnt,label)

def isTestTable(table):
    bRet = False

    if len(table.rows) > 1:
        for r, rowiteration in enumerate(table.rows):
            for c, cells in enumerate(rowiteration.cells):
                if cells.text == "IOP Test Case Reference":
                    bRet= True
                    break

    return bRet
        
            
