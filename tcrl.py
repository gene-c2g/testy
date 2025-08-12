##################################################
## testy.py
##################################################
## Author: clenahan@cloud2gnd.com
## Copyright: Copyright 2024
## Version: 1.0
##################################################

from bcolors import bcolors
from wordhelper import getAcceptedText, addNewline
from excelhelper import setHeaderCell

from docx import Document
from docx.enum.section import WD_ORIENT,WD_SECTION
import openpyxl
from openpyxl.formatting.rule import CellIsRule, FormulaRule

import os
import re


def dumpTcrl(filename,naturalSort):
    if  not os.path.exists(filename):
        print(bcolors.FAIL + "File '%s' does not exist" % filename + bcolors.ENDC)
        return
    
    doc = Document(filename)

    tcidlist = []

    parseTCIDs(doc, lambda x: tcidlist.append(x))
    if not naturalSort:
        tcidlist.sort()

    outputListXlsx(tcidlist,filename)

    outputListDocx(tcidlist,filename)

def extractTCID(text):
     reg = "/*\-[IC]"
     match= re.search(reg,text)
     return match.string if match else None


    
def parseTCIDs(doc,delegate):
    for para in doc.paragraphs:
        if "Heading 9" in para.style.name or "Heading 8" in para.style.name :
            match= extractTCID(getAcceptedText(para))
            if match:
                delegate(match)

    for table in doc.tables:
        for r, rowiteration in enumerate(table.rows):
            for c, cells in enumerate(rowiteration.cells):
                bFound = False
                if len(cells.paragraphs) >= 1:
                    para = cells.paragraphs[0]
                    text = getAcceptedText(para)
                    match= extractTCID(text)
                    if match:
                        if "Heading 9" in para.style.name or "Heading 8" in para.style.name :
                            bFound = True
                    if bFound:
                        delegate(match)
              

def outputListXlsx(tcidlist,filename):
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.sheet_view.zoomScale = 140
    setHeaderCell(sheet,1,"TCID",100)
    setHeaderCell(sheet,2,"TCID",30)
    setHeaderCell(sheet,3,"Description",100)

    nRow=2
    for rawtcid in tcidlist:
        sheet.cell(row = nRow, column =1).value = rawtcid
        tcid = rawtcid.split(" ")
        description = " ".join(tcid[1:]).replace("[","").replace("]","")
        sheet.cell(row = nRow, column =2).value = tcid[0]
        sheet.cell(row = nRow, column =3).value = description
        nRow += 1

    (root,ext) = os.path.splitext(filename)
    newfile = root+"_TCRL.xlsx"
    wb.save(newfile)
    print("Created %s" % newfile)


def passCount(tcid):
    basetcid = tcid.split(" ")[0].split("/")
    root = basetcid[-1][0:2]
    return 2 if root == "BI" else 3

def isOptionalTest(tcid):
    return True if "GGIT" in tcid else False

def outputListDocx(tcidlist,filename):
    doc = Document()

    table= doc.add_table(rows=0,cols=4)
    addTCRLTable(table,tcidlist)

    addNewline(doc)

    table= doc.add_table(rows=0,cols=3)
    addTCMTTable(table,tcidlist)

    (root,ext) = os.path.splitext(filename)
    newfile =root+"_TCRL.docx"
    doc.save(newfile)
    print("Created %s" % newfile)

def addTCMTTable(table,tcidlist):
    for rawtcid in tcidlist:
        row = table.add_row().cells 
        regex= "(.*) \[(.*)\]"
        match= re.search(regex,rawtcid)
        if match:
            row[2].text = match.group(1)
            row[1].text = match.group(2)

def addTCRLTable(table, tcidlist):
    nRow = 1
    for rawtcid in tcidlist:
        row = table.add_row().cells 
        row[0].text = rawtcid
        row[1].text = "XXX-" + ("%d" % nRow).zfill(3)
        if isOptionalTest(rawtcid):
            row[2].text = "0"
            row[3].text = "Generic"
        else:
            row[2].text = str(passCount(rawtcid))
            row[3].text = "Default"
        nRow+=1
